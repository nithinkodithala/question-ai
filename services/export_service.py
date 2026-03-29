from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import io
import textwrap
import os
import re


def _remove_trailing_empty_paragraphs(doc):
    """Scan and remove empty paragraphs from the bottom of the template."""
    while len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        # If it has no text and no tables/images/breaks, try to remove it
        if not last_para.text.strip():
            # In python-docx, deleting a paragraph is tricky. 
            # We clear it to minimize space, or we can use the internal element removal.
            p = last_para._element
            p.getparent().remove(p)
            p._p = p._element = None
        else:
            break

SECTION_DESCRIPTIONS_40M = {
    r'^SECTION\s+A\b': "(Section-A consists of 10 questions 1 mark each)",
    r'^SECTION\s+B\b': "(Section-B consists of 4 questions 2 marks each)",
    r'^SECTION\s+C\b': "(Section-C consists of 3 questions 3 marks each)",
    r'^SECTION\s+D\b': "(Section-D consists of 1 question 5 marks each)",
    r'^SECTION\s+E\b': "(Case study questions)",
}

def _inject_section_descriptions(text: str) -> str:
    """Ensure section descriptions are present after section headers for 40M papers."""
    lines = text.split('\n')
    new_lines = []
    for i, line in enumerate(lines):
        new_lines.append(line)
        stripped = line.strip()
        for pattern, description in SECTION_DESCRIPTIONS_40M.items():
            if re.match(pattern, stripped, re.IGNORECASE):
                next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
                if description not in next_line:
                    new_lines.append(description)
                break
    return '\n'.join(new_lines)

def _strip_question_marks(text: str) -> str:
    """Remove individual per-question marks like (1), (2 M), (1 mark) anywhere in the text."""
    # Handles: (1 mark), (1 M), (1), (2 marks), also with trailing ? 
    return re.sub(r'\s*\(\s*\d+\s*(?:marks?|M|m)?\s*\)\s*\??', '', text)

def _format_question_paper(doc, text: str, max_marks: str = "40"):
    """Helper to add formatted questions to a document."""
    # Guaranteed pre-processing: strip marks + inject section descriptions
    text = _strip_question_marks(text)
    if str(max_marks) == "40":
        text = _inject_section_descriptions(text)

    lines = text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
        para = doc.add_paragraph()
        
        # Detect section header
        if re.search(r'^SECTION\s+[A-Z]', clean_line, re.IGNORECASE):
            run = para.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(13)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(2)
            
            # Inject section description immediately for 40M paper
            if str(max_marks) == "40":
                sec_match = re.search(r'SECTION\s+([A-E])', clean_line, re.IGNORECASE)
                if sec_match:
                    sec_letter = sec_match.group(1).upper()
                    desc_map = {
                        "A": "(Section-A consists of 10 questions 1 mark each)",
                        "B": "(Section-B consists of 4 questions 2 marks each)",
                        "C": "(Section-C consists of 3 questions 3 marks each)",
                        "D": "(Section-D consists of 1 question 5 marks each)",
                        "E": "(Case study questions)",
                    }
                    if sec_letter in desc_map:
                        desc_para = doc.add_paragraph()
                        desc_run = desc_para.add_run(desc_map[sec_letter])
                        desc_run.bold = True
                        desc_run.font.size = Pt(10)
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        desc_para.paragraph_format.space_before = Pt(0)
                        desc_para.paragraph_format.space_after = Pt(6)

        elif clean_line.startswith('(') and clean_line.endswith(')'):
            # Skip if already-injected section description to avoid duplicates
            is_sec_desc = any(k in clean_line for k in [
                "Section-A consists", "Section-B consists",
                "Section-C consists", "Section-D consists", "Case study"
            ])
            if is_sec_desc:
                # Remove the empty paragraph created at top of loop
                p = para._element
                p.getparent().remove(p)
            else:
                run = para.add_run(clean_line)
                run.italic = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(6)

        else:
            # Handle numbered questions with vertical alignment
            q_match = re.match(r'^(\d+)\.\s*(.*)', clean_line)
            if q_match:
                q_num = q_match.group(1)
                q_text = q_match.group(2)
                
                # Clear existing content if any (para was created in the loop)
                # Then add formatted text with tab
                para.paragraph_format.tab_stops.add_tab_stop(Inches(0.4))
                # Set a hanging indent for multi-line questions
                para.paragraph_format.left_indent = Inches(0.4)
                para.paragraph_format.first_line_indent = Inches(-0.4)
                
                # Detect marks distribution pattern (e.g. (5 M))
                marks_match = re.search(r'\(([^)]+\d+\s*M)\)\s*$', q_text)
                if marks_match:
                    marks_text = marks_match.group(0)
                    main_q_text = q_text.replace(marks_text, "").strip()
                    para.add_run(f"{q_num}.\t{main_q_text}")
                    para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
                    para.add_run("\t")
                    m_run = para.add_run(marks_text)
                    m_run.bold = True
                else:
                    para.add_run(f"{q_num}.\t{q_text}")
                continue

            # Handle MCQ options with alignment (supporting multiple on one line)
            if re.match(r'^A\.', clean_line):
                # Format horizontal options A. ... B. ... C. ... D. ...
                # Space them out with tabs
                formatted_opts = clean_line.replace(" B.", "\tB.").replace(" C.", "\tC.").replace(" D.", "\tD.")
                para.paragraph_format.left_indent = Inches(0.4)
                
                # Set tab stops for options to align them in columns
                para.paragraph_format.tab_stops.add_tab_stop(Inches(2.0))
                para.paragraph_format.tab_stops.add_tab_stop(Inches(3.6))
                para.paragraph_format.tab_stops.add_tab_stop(Inches(5.2))
                
                para.add_run(formatted_opts)
                continue
                
            para.add_run(clean_line)

def generate_docx_from_text(text: str, school_name: str = "GENIUS HIGH SCHOOL :: BHONGIR", test_name: str = "Periodic Test - I", subject: str = "Science", class_name: str = "VII", time_limit: str = "90mins", max_marks: str = "40") -> bytes:
    """Generate Word doc with branded header and formatted questions."""
    print(f"DEBUG: Building document for {school_name} - {test_name}")
    buffer = io.BytesIO()
    doc = Document()
    
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def zero_spacing(para):
        """Remove all before/after spacing from a paragraph."""
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(13)

    # 1. Official Branded Header (Table-based alignment)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.columns[0].width = Inches(1.0)
    header_table.columns[1].width = Inches(4.5)
    header_table.columns[2].width = Inches(1.0)
    
    cells = header_table.rows[0].cells
    
    # Left: Logo (smaller)
    logo_path = os.path.join("static", "logo.png")
    if os.path.exists(logo_path):
        p_logo = cells[0].add_paragraph()
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(0.7))
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        zero_spacing(p_logo)
    
    # Middle: School Name (compact)
    p_school = cells[1].add_paragraph()
    run_name = p_school.add_run(school_name.upper())
    run_name.bold = True
    run_name.font.size = Pt(12)
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zero_spacing(p_school)
    
    # Test Title on next line (compact)
    p_test = cells[1].add_paragraph()
    run_test = p_test.add_run(test_name.upper())
    run_test.bold = True
    run_test.font.underline = True
    run_test.font.size = Pt(11)
    p_test.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zero_spacing(p_test)
    
    # Details Row: Class | Marks on one row, Subject | Time on second row
    row_details = doc.add_table(rows=2, cols=2)
    
    for row in row_details.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                zero_spacing(para)

    # Row 1: Class | Marks
    c_p1 = row_details.rows[0].cells[0].add_paragraph()
    c_run = c_p1.add_run(f"Class: {class_name}")
    c_run.bold = True
    c_run.font.size = Pt(10)
    zero_spacing(c_p1)
    
    m_p1 = row_details.rows[0].cells[1].add_paragraph()
    m_run = m_p1.add_run(f"Max. Marks: {max_marks}")
    m_run.bold = True
    m_run.font.size = Pt(10)
    m_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    zero_spacing(m_p1)

    # Row 2: Subject | Time
    s_p1 = row_details.rows[1].cells[0].add_paragraph()
    s_run = s_p1.add_run(f"Subject: {subject}")
    s_run.bold = True
    s_run.font.size = Pt(10)
    zero_spacing(s_p1)
    
    t_p1 = row_details.rows[1].cells[1].add_paragraph()
    t_run = t_p1.add_run(f"Time: {time_limit}")
    t_run.bold = True
    t_run.font.size = Pt(10)
    t_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    zero_spacing(t_p1)
    
    # Thin separator line
    sep = doc.add_paragraph("_" * 90)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zero_spacing(sep)
    
    # 2. General Instructions for 40/80 Marks Paper
    if str(max_marks) == "40":
        instructions_para = doc.add_paragraph()
        run_inst = instructions_para.add_run("General Instructions:")
        run_inst.bold = True
        run_inst.font.underline = True
        
        instructions = [
            "i.This question paper consists of 20 questions in 5 sections.",
            "ii.All questions are compulsory. However, an internal choice is provided in some questions. A student is expected to attempt only one of these questions.",
            "iii.Section A consists of 10 MCQ's type questions carrying 1 mark each.",
            "iv.Section B consists of 4 Very Short questions carrying 02 marks each. Answers to these questions should be in the range of 30 to 50 words.",
            "v.Section C consists of 3 Short Answer type questions carrying 03 marks each. Answers to these Questions should in the range of 50 to 80 words",
            "vi.Section D consists of 1 Long Answer type questions carrying 05 marks each. Answer to these Questions should be in the range of 80 to 120 words.",
            "vii.Section E consists of 2 source-based/case-based units of assessment of 04 marks each with sub-parts."
        ]
        
        for inst in instructions:
            p = doc.add_paragraph()
            run = p.add_run(inst)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(2)
            run.font.size = Pt(10)
            
        doc.add_paragraph("_" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    elif str(max_marks) == "80":
        instructions_para = doc.add_paragraph()
        run_inst = instructions_para.add_run("General Instructions:")
        run_inst.bold = True
        run_inst.font.underline = True
        
        instructions = [
            "1. This Question Paper has 5 Sections A-E.",
            "2. Section A has 20 MCQs carrying 1 mark each.",
            "3. Section B has 5 questions carrying 02 marks each.",
            "4. Section C has 6 questions carrying 03 marks each.",
            "5. Section D has 4 questions carrying 05 marks each.",
            "6. Section E has 3 case based integrated units of assessment (04 marks each) with sub-parts of the values of 1, 1 and 2 marks each respectively.",
            "7. All Questions are compulsory. However, an internal choice in 2 Qs of 5 marks, 2 Qs of 3 marks and 2 Questions of 2 marks has been provided. An internal choice has been provided in the 2marks questions of Section E.",
            "8. Draw neat figures wherever required. Take π =22/7 wherever required if not stated."
        ]
        
        for inst in instructions:
            p = doc.add_paragraph()
            run = p.add_run(inst)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(2)
            run.font.size = Pt(10)
            
        doc.add_paragraph("_" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Content
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    _format_question_paper(doc, text, max_marks=str(max_marks))
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_docx_with_template(text: str, template_bytes: bytes, subject: str = "", class_name: str = "") -> bytes:
    """DOCX template support."""
    buffer = io.BytesIO()
    doc = Document(io.BytesIO(template_bytes))
    placeholder_found = False
    # Handle paragraphs
    for paragraph in doc.paragraphs:
        if '{{QUESTIONS}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{QUESTIONS}}', text)
            placeholder_found = True
        if '{{SUBJECT}}' in paragraph.text and subject:
            paragraph.text = paragraph.text.replace('{{SUBJECT}}', subject)
        if '{{CLASS}}' in paragraph.text and class_name:
            paragraph.text = paragraph.text.replace('{{CLASS}}', class_name)
            
    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{QUESTIONS}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{QUESTIONS}}', text)
                        placeholder_found = True
                    if '{{SUBJECT}}' in paragraph.text and subject:
                        paragraph.text = paragraph.text.replace('{{SUBJECT}}', subject)
                    if '{{CLASS}}' in paragraph.text and class_name:
                        paragraph.text = paragraph.text.replace('{{CLASS}}', class_name)
    if not placeholder_found:
        # Clean up the end of the template to avoid extra pages/space
        _remove_trailing_empty_paragraphs(doc)
        
        # Start immediately below the last template content
        p = doc.add_paragraph()
        run = p.add_run('QUESTIONS')
        run.bold = True
        run.font.underline = True
        run.font.size = Pt(14)
        
        _format_question_paper(doc, text)
        
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
